"""DOCX exporter using python-docx."""

import asyncio
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


async def export_to_docx(content: Dict[str, Any], name: str) -> str:
    doc = Document()

    title = doc.add_heading(name, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    summary = content.get("summary", "")
    if summary:
        p = doc.add_paragraph(summary)
        p.style = "Quote"

    doc.add_heading("项目经历", 1)
    for project in content.get("projects", []):
        doc.add_heading(project.get("name", ""), 2)
        role = project.get("role", "")
        start_date = project.get("start_date", "")
        end_date = project.get("end_date", "")
        if role or start_date or end_date:
            p = doc.add_paragraph()
            p.add_run("角色: ").bold = True
            p.add_run(role)
            p = doc.add_paragraph()
            p.add_run("时间: ").bold = True
            p.add_run(f"{start_date} - {end_date}")

        star_fields = [
            ("背景", "situation"),
            ("任务", "task"),
            ("行动", "action"),
            ("成果", "result"),
        ]
        has_star = any(project.get(k) for _, k in star_fields)
        if has_star:
            for star_label, star_key in star_fields:
                val = project.get(star_key, "")
                if val:
                    p = doc.add_paragraph()
                    p.add_run(f"{star_label}: ").bold = True
                    p.add_run(val)
        elif project.get("description"):
            doc.add_paragraph(project["description"])

    doc.add_heading("技能", 1)
    skills = content.get("skills", [])
    if skills:
        p = doc.add_paragraph(", ".join(skills))

    experience = content.get("experience", [])
    if experience:
        doc.add_heading("工作经历", 1)
        for exp in experience:
            if isinstance(exp, str):
                doc.add_paragraph(exp)
            elif isinstance(exp, dict):
                p = doc.add_paragraph()
                p.add_run(f"{exp.get('company', exp.get('name', ''))}").bold = True
                if exp.get('role') or exp.get('position'):
                    p.add_run(f" | {exp.get('role', exp.get('position', ''))}")
                if exp.get('period') or exp.get('time'):
                    p.add_run(f" | {exp.get('period', exp.get('time', ''))}")
                for key in ['description', 'highlights']:
                    if exp.get(key):
                        items = exp[key] if isinstance(exp[key], list) else [exp[key]]
                        for item in items:
                            doc.add_paragraph(str(item))

    education = content.get("education", [])
    if education:
        doc.add_heading("教育背景", 1)
        for edu in education:
            if isinstance(edu, str):
                doc.add_paragraph(edu)
            elif isinstance(edu, dict):
                p = doc.add_paragraph()
                p.add_run(f"{edu.get('school', edu.get('name', ''))}").bold = True
                if edu.get('major'):
                    p.add_run(f" | {edu.get('major', '')}")
                if edu.get('degree'):
                    p.add_run(f" | {edu.get('degree', '')}")

    output_dir = Path("data/exports")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{name}.docx"

    loop = asyncio.get_running_loop()
    await loop.run_in_executor(None, lambda: doc.save(str(output_path)))
    return str(output_path)
